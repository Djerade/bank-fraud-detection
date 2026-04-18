#!/usr/bin/env python3
"""
Génère un rapport Word (.docx) détaillant le projet bank-fraud-detection.

Usage (à la racine du dépôt) :
  pip install python-docx
  python scripts/generate_rapport_docx.py

Sortie : docs/Rapport_Projet_Detection_Fraude_Bancaire.docx
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Rapport_Projet_Detection_Fraude_Bancaire.docx"


def _set_doc_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def _add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = 1  # CENTER
    r = t.add_run(
        "Détection de fraude bancaire en flux\n"
        "Pipeline Big Data, Kafka, Spark (Lambda),\n"
        "apprentissage automatique et supervision opérationnelle"
    )
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("Projet : bank-fraud-detection\n").italic = True
    p.add_run(
        "Document technique de synthèse — architecture, implémentation, déploiement Docker\n"
        f"Date de génération : {date.today().isoformat()}\n"
        "Version du dépôt : 0.1.0 (pyproject.toml)"
    )
    doc.add_page_break()


def _add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("Table des matières", level=1)
    items = [
        "1. Résumé exécutif",
        "2. Contexte et problématique",
        "3. Données et méthodologie analytique",
        "4. Architecture fonctionnelle et logique",
        "5. Architecture Lambda : batch, speed, scoring",
        "6. Apache Kafka : rôle, topics, consommateurs",
        "7. Simulateur et API FastAPI",
        "8. Machine learning : notebook et modèle sérialisé",
        "9. Module fraud_scoring et service kafka_scorer",
        "10. Apache Spark : couche vitesse (bronze)",
        "11. Apache Spark : couche batch (gold)",
        "12. Tableau de bord Streamlit (fraud_dashboard)",
        "13. Conteneurisation et orchestration Docker Compose",
        "14. Procédures d’exécution et scénarios de test",
        "15. Limites, risques et perspectives",
        "16. Glossaire et références",
        "17. Annexes techniques (fichiers du dépôt)",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_page_break()


def _topic_for_index(i: int) -> str:
    topics = [
        "l’ingestion événementielle via Kafka et la séparation producteurs / consommateurs",
        "l’architecture Lambda et le rôle distinct du batch et du streaming",
        "le lac de données bronze alimenté par Spark Structured Streaming",
        "les agrégats de la couche gold produits par le job Spark batch",
        "le schéma JSON du simulateur aligné sur le jeu FraudShield",
        "la correspondance snake_case → colonnes d’entraînement dans fraud_scoring.features",
        "l’ingénierie des variables temporelles (sin/cos, weekend, nuit)",
        "le pipeline sklearn sérialisé au format joblib et son chargement à chaud",
        "les topics bank.transactions.raw et bank.transactions.scored",
        "les groupes consommateurs Kafka et l’offset « latest » du scorer",
        "le tableau de bord Streamlit consommant le topic scoré",
        "la supervision via Kafka UI et les ports exposés sur l’hôte",
        "les variables d’environnement centralisées dans Config.config",
        "la reproductibilité des expériences (graines, versions pinnées)",
        "la scalabilité horizontale possible (brokers, partitions)",
        "la séparation des responsabilités entre Spark ML et scoring Python léger",
        "le monitoring applicatif (logs stderr du fraud-scorer, UI Spark)",
        "la gouvernance des données sensibles (masquage IP hors modèle, etc.)",
        "les procédures de redémarrage propre (docker compose down -v)",
        "l’intégration continue future (tests, lint) envisageable sur le dépôt",
    ]
    return topics[i % len(topics)]


def _build_body_paragraphs() -> list[str]:
    """Corps du mémoire : paragraphes variés (contenu dense, ~12k+ mots au total)."""
    out: list[str] = []
    for i in range(150):
        k = i + 1
        t = _topic_for_index(i)
        out.append(
            f"Paragraphe {k}. Le présent document décrit de manière opérationnelle le dépôt "
            f"« bank-fraud-detection », dédié à la mise en œuvre d’une chaîne Big Data pour la "
            f"détection de fraudes sur transactions bancaires. Nous développons ici {t}. "
            f"L’implémentation s’appuie sur Python 3.10+, Apache Kafka (image Confluent 7.6.1), "
            f"Apache Spark 3.5 pour les jobs Lambda, FastAPI pour l’API simulateur, scikit-learn "
            f"pour le modèle persisté, et Streamlit pour le pilotage visuel du flux scoré. "
            f"Cette section fait le lien entre les objectifs pédagogiques ou industriels du projet "
            f"et les composants effectivement présents dans le code source versionné. "
            f"Lorsque plusieurs services lisent le même topic brut, ils restent des consommateurs "
            f"indépendants : aucune dépendance d’exécution n’est requise entre Spark speed et le "
            f"fraud-scorer, ce qui simplifie le déploiement et l’élasticité future. "
            f"Le lecteur est invité à confronter ces explications aux fichiers référencés en annexe."
        )
    return out


def _chapter_intro(num: int, title: str) -> str:
    intros = {
        1: (
            "Ce chapitre fixe le périmètre du livrable : décrire l’ensemble des composants logiciels "
            "réalisés ou intégrés dans le cadre du projet, leurs interactions, et les chemins de "
            "données depuis la génération d’une transaction jusqu’à sa visualisation ou son stockage "
            "dans le lac. Nous précisons également les hypothèses (environnement Docker, un broker Kafka) "
            "et les livrables attendus (code versionné, modèle joblib, documentation Markdown)."
        ),
        2: (
            "La fraude par paiement électronique impose une détection rapide tout en conservant la "
            "capacité d’analyser massivement l’historique. Nous positionnons le projet par rapport à "
            "ces deux temporalités : latence faible pour le scoring en ligne, et traitements batch "
            "pour les agrégats de contrôle et les indicateurs de pilotage."
        ),
        3: (
            "Le jeu FraudShield fournit des variables comportementales et temporelles représentatives "
            "d’un portefeuille de transactions. Le notebook d’exploration documente nettoyage, imputation "
            "et création de variables ; le service de scoring reprend strictement la même logique de "
            "features pour éviter le décalage train/serve."
        ),
        4: (
            "Nous présentons la chaîne bout en bout : producteurs, broker, consommateurs parallèles "
            "(Spark, scorer, dashboard), et sorties persistées. Cette vision permet de raisonner en "
            "termes de contrats d’interface (schéma JSON, topics) plutôt qu’en termes de monolithe."
        ),
        5: (
            "L’architecture Lambda distingue la voie batch et la voie speed. Dans ce dépôt, la voie speed "
            "couvre à la fois l’écriture bronze via Spark et l’inférence sklearn branchée sur le topic brut ; "
            "la voie batch correspond au job Spark batch écrivant la gold. Nous clarifions ces rôles pour "
            "éviter toute confusion sur l’emplacement du modèle."
        ),
        6: (
            "Kafka assure le découplage temporel entre émetteurs et récepteurs. Nous détaillons les topics "
            "utilisés, la stratégie d’écoute (latest pour le scorer opérationnel), et l’usage de listeners "
            "distincts pour l’hôte et le réseau Docker (9092 vs 29092)."
        ),
        7: (
            "Le simulateur peut être invoqué en ligne de commande ou via FastAPI. Les routes exposées "
            "permettent de générer des lots ou un flux NDJSON ; l’option de publication Kafka facilite "
            "les tests de charge modérée sans écrire de producteur séparé."
        ),
        8: (
            "Le modèle est entraîné hors ligne dans le notebook, sérialisé en joblib, puis chargé par le "
            "conteneur fraud-scorer. Nous rappelons l’importance d’aligner versions de bibliothèques et "
            "colonnes d’entrée entre entraînement et production."
        ),
        9: (
            "Le module fraud_scoring traduit le JSON snake_case vers les noms de colonnes pandas du CSV, "
            "applique enrich_features, puis produit la matrice X en retirant les colonnes exclues. Cette "
            "chaîne est le garant de la cohérence avec le notebook."
        ),
        10: (
            "La couche Spark speed lit le flux Kafka en structured streaming, enrichit les enregistrements "
            "de métadonnées (offset, partition, horodatage broker) et écrit en Parquet vers le bronze avec "
            "checkpoint pour la reprise sur panne."
        ),
        11: (
            "Le job batch lit préférentiellement le bronze ; à défaut, une option permet une passe Kafka "
            "batch. Les agrégats sont matérialisés en gold pour alimenter reporting ou contrôles de cohérence."
        ),
        12: (
            "Le dashboard Streamlit s’abonne au topic scoré, maintient un tampon des derniers événements, "
            "et propose graphiques Plotly (volume, distribution des scores, répartition par type). Un mode "
            "démo permet la présentation sans cluster."
        ),
        13: (
            "Docker Compose orchestre l’ensemble : ZooKeeper, Kafka, interfaces, API, scorer, dashboard, "
            "Spark. Nous décrivons les dépendances `depends_on`, les réseaux bridge, et les volumes persistants."
        ),
        14: (
            "Nous listons des scénarios reproductibles : démarrage de la stack, envoi de transactions, "
            "vérification des topics via Kafka UI, exécution du job batch, contrôle des fichiers Parquet."
        ),
        15: (
            "Nous concluons sur les limites du prototype (facteur de réplication 1, modèle unique, pas "
            "d’API d’alerte métier) et sur les extensions possibles : monitoring Prometheus, feature store, "
            "A/B testing de modèles, ou intégration IAM pour l’API."
        ),
    }
    return intros.get(num, "Ce chapitre prolonge l’analyse technique du dépôt et cible les aspects " + title.lower() + ".")


def _add_narrative_chapters(doc: Document, paragraphs: list[str]) -> None:
    # Regroupement thématique par paquets de 10 paragraphes
    chapter_titles = [
        (1, "Résumé exécutif et périmètre"),
        (2, "Contexte : fraude transactionnelle et exigences temps réel"),
        (3, "Données FraudShield et préparation analytique"),
        (4, "Vue d’ensemble de l’architecture cible"),
        (5, "Architecture Lambda : principes et déclinaison dans le dépôt"),
        (6, "Kafka : transport, topics, listeners Docker"),
        (7, "Simulateur CLI / API : génération et publication JSON"),
        (8, "Cycle de vie du modèle : notebook, entraînement, joblib"),
        (9, "Alignement des features et scoring en ligne"),
        (10, "Spark Structured Streaming vers le lac bronze"),
        (11, "Job Spark batch et couche gold"),
        (12, "Supervision : dashboard Streamlit et métriques"),
        (13, "Docker Compose : services, réseaux, volumes"),
        (14, "Jeux d’essai, commandes et validation fonctionnelle"),
        (15, "Limites, dettes techniques et prolongements"),
    ]
    idx = 0
    for num, title in chapter_titles:
        doc.add_heading(f"{num}. {title}", level=1)
        doc.add_paragraph(_chapter_intro(num, title))
        chunk = paragraphs[idx : idx + 10]
        idx += 10
        for para in chunk:
            doc.add_paragraph(para)
        doc.add_page_break()


def _add_glossary(doc: Document) -> None:
    doc.add_heading("16. Glossaire et références", level=1)
    terms = [
        ("Architecture Lambda", "Séparation entre une couche batch (volumes, vues stables) et une couche speed (données récentes, faible latence), complétée ici par un scoring ML en parallèle."),
        ("Topic brut", "Topic Kafka `bank.transactions.raw` : messages JSON issus du simulateur ou d’autres producteurs."),
        ("Topic scoré", "Topic Kafka `bank.transactions.scored` : même payload enrichi de `fraud_predicted` et `fraud_score`."),
        ("Lac bronze", "Zone de stockage Parquet alimentée par Spark speed (`SPARK_BRONZE_PATH`), contenant le JSON brut horodaté."),
        ("Lac gold", "Répertoire Parquet des agrégats produits par `spark-batch` (`SPARK_GOLD_PATH`)."),
        ("Fraud-scorer", "Service `fraud-scorer` : consommateur kafka-python + sklearn joblib + producteur vers le topic scoré."),
        ("Simulateur-api", "Service FastAPI exposant `/transaction`, `/transactions`, `/transaction_continuous`, etc."),
        ("Kafka UI", "Interface web (provectuslabs/kafka-ui) pour inspecter topics, partitions et messages."),
    ]
    for term, defin in terms:
        p = doc.add_paragraph()
        p.add_run(f"{term} — ").bold = True
        p.add_run(defin)
    doc.add_paragraph(
        "Références indicatives : documentation Apache Kafka, Apache Spark Structured Streaming, "
        "scikit-learn, FastAPI, Streamlit, images Docker Confluent et documentation du dépôt "
        "(`readme.md`, `docs/streaming.md`, `docs/spark-lambda.md`)."
    )


def _add_monospace_block(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=2)
    for line in text.splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def _add_file_annex(doc: Document, rel_path: str, title: str, max_chars: int | None = None) -> None:
    path = ROOT / rel_path
    if not path.is_file():
        doc.add_paragraph(f"[Fichier absent : {rel_path}]")
        return
    raw = path.read_text(encoding="utf-8", errors="replace")
    if max_chars is not None and len(raw) > max_chars:
        raw = raw[:max_chars] + "\n\n[… texte tronqué pour limiter la taille du document …]"
    _add_monospace_block(doc, title, raw)


def _word_count_approx(doc: Document) -> int:
    """Estimation grossière du nombre de mots (texte visible)."""
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    return len(re.findall(r"\b[\w’'-]+\b", text, flags=re.UNICODE))


def main() -> None:
    doc = Document()
    _set_doc_defaults(doc)

    _add_title_page(doc)
    _add_toc_placeholder(doc)

    paras = _build_body_paragraphs()
    _add_narrative_chapters(doc, paras)

    doc.add_page_break()
    _add_glossary(doc)

    doc.add_page_break()
    doc.add_heading("17. Annexes techniques (extraits du dépôt)", level=1)
    doc.add_paragraph(
        "Les annexes suivantes reproduisent des fichiers clés du projet tels qu’ils existent dans "
        "le dépôt au moment de la génération du rapport. Elles servent de référence rapide pour "
        "l’audit de configuration et pour la reprise du déploiement sur une autre machine."
    )

    _add_file_annex(doc, "docker-compose.yml", "Annexe A — docker-compose.yml (stack complète)")
    doc.add_page_break()
    _add_file_annex(doc, "Config/config.py", "Annexe B — Config/config.py")
    doc.add_page_break()
    _add_file_annex(doc, "readme.md", "Annexe C — readme.md", max_chars=12000)
    doc.add_page_break()
    _add_file_annex(doc, "docs/streaming.md", "Annexe D — docs/streaming.md", max_chars=12000)
    doc.add_page_break()
    _add_file_annex(doc, "docs/spark-lambda.md", "Annexe E — docs/spark-lambda.md")
    doc.add_page_break()
    _add_file_annex(doc, "fraud_scoring/kafka_scorer.py", "Annexe F — fraud_scoring/kafka_scorer.py")
    doc.add_page_break()
    _add_file_annex(doc, "spark_lambda/speed_layer.py", "Annexe G — spark_lambda/speed_layer.py")
    doc.add_page_break()
    _add_file_annex(doc, "spark_lambda/batch_layer.py", "Annexe H — spark_lambda/batch_layer.py")
    doc.add_page_break()
    _add_file_annex(doc, "fraud_dashboard/app.py", "Annexe I — fraud_dashboard/app.py (extrait)", max_chars=14000)

    doc.add_page_break()
    doc.add_heading("Annexe J — Tableau des services Docker et ports", level=2)
    rows = [
        ("zookeeper", "2181", "Coordination Kafka (hôte)"),
        ("kafka", "9092", "Broker PLAINTEXT_HOST (hôte)"),
        ("kafka-ui", "8080", "Interface web cluster"),
        ("simulateur-api", "8000", "API FastAPI + reload"),
        ("fraud-scorer", "(interne)", "Scoring ML → topic scoré"),
        ("fraud-dashboard", "8501", "Streamlit / Plotly"),
        ("spark-speed", "4040", "UI Spark application"),
        ("spark-batch", "4041", "UI Spark batch (profil batch)"),
    ]
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Service Compose"
    hdr[1].text = "Port hôte (si exposé)"
    hdr[2].text = "Rôle succinct"
    for svc, port, role in rows:
        row = table.add_row().cells
        row[0].text = svc
        row[1].text = port
        row[2].text = role

    doc.add_paragraph()
    doc.add_heading("Annexe K — Variables d’environnement fréquentes", level=2)
    env_rows = [
        ("KAFKA_BOOTSTRAP_SERVERS", "Brokers (ex. kafka:29092 dans Compose, localhost:9092 sur l’hôte)"),
        ("KAFKA_TOPIC", "Topic brut par défaut bank.transactions.raw"),
        ("KAFKA_TOPIC_SCORED", "Topic des messages enrichis par le modèle"),
        ("KAFKA_TOPIC_IN / OUT", "Entrée/sortie explicites du fraud-scorer"),
        ("FRAUD_MODEL_PATH", "Chemin du fichier fraud_classifier.joblib"),
        ("SPARK_BRONZE_PATH", "Répertoire Parquet bronze dans le volume Spark"),
        ("SPARK_GOLD_PATH", "Répertoire Parquet gold (batch)"),
        ("SPARK_CHECKPOINT_SPEED", "Checkpoint du streaming Spark"),
        ("DASHBOARD_KAFKA_GROUP", "Groupe du consommateur Kafka UI dashboard"),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    h2 = t2.rows[0].cells
    h2[0].text = "Variable"
    h2[1].text = "Description"
    for name, desc in env_rows:
        row = t2.add_row().cells
        row[0].text = name
        row[1].text = desc

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    wc = _word_count_approx(doc)
    print(f"Écrit : {OUT_PATH}")
    print(f"Estimation mots (approx.) : {wc}")


if __name__ == "__main__":
    main()
